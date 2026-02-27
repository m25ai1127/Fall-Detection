"""
Generate Study Notes as a Word Document
"""
import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent.parent))

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ---- Styles ----
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h

def add_para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            table.rows[r_idx + 1].cells[c_idx].text = str(cell_text)
    doc.add_paragraph()  # spacing


# ============================================================
# TITLE
# ============================================================
title = doc.add_heading('Fall Detection Project - Complete Study Notes', level=0)
for run in title.runs:
    run.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)
add_para('Everything you need to know, topic by topic', bold=True)
doc.add_paragraph()

# ============================================================
# 1. Python Basics
# ============================================================
add_heading('1. Python Basics Used')

add_heading('import os', level=2)
add_bullet('os = Operating System module. Lets Python talk to your computer.')
add_bullet('os.path.join("data", "raw") -> creates path data\\raw')
add_bullet('os.makedirs("results") -> creates a folder')
add_bullet('os.path.exists("file.txt") -> checks if file exists (True/False)')

add_heading('from pathlib import Path', level=2)
add_bullet('Modern way to handle file paths (cleaner than os.path)')
add_bullet('Path("data") / "raw" / "video.mp4" -> creates data\\raw\\video.mp4')
add_bullet('Path(__file__).parent -> folder where the current file lives')
add_bullet('All paths are relative to project root, so it works no matter where you move the folder')

# ============================================================
# 2. MiDaS
# ============================================================
add_heading('2. MiDaS Depth Estimation')

add_heading('Three Model Sizes', level=2)
add_table(
    ['Model', 'Speed', 'Accuracy', 'Use Case'],
    [
        ['MiDaS_small (ours)', 'Fast', 'Good', 'Real-time on CPU'],
        ['DPT_Hybrid', 'Medium', 'Better', 'ResNet + Transformer combined'],
        ['DPT_Large', 'Slow', 'Best', 'Full Transformer, needs GPU'],
    ]
)

add_heading('DPT = Dense Prediction Transformer', level=2)
add_bullet('Dense', ' = predicts a value for EVERY pixel')
add_bullet('Prediction', ' = outputs depth values')
add_bullet('Transformer', ' = same architecture behind ChatGPT, but for images')
add_bullet('Hybrid', ' = CNN (ResNet) at bottom + Transformer at top')

add_heading('Processing Time on CPU', level=2)
add_table(
    ['Model', 'Per Frame', 'Live Camera FPS'],
    [
        ['MiDaS_small', '~0.05s', '~5-8 FPS'],
        ['DPT_Hybrid', '~0.3s', '~1-2 FPS'],
        ['DPT_Large', '~0.8s', '<1 FPS (unusable)'],
    ]
)

add_heading('Device Auto-Detection', level=2)
add_para('MIDAS_DEVICE = "cuda" if GPU available, else "cpu"')
add_bullet('Checks if NVIDIA GPU exists -> uses GPU (fast), otherwise CPU')

# ============================================================
# 3. Pre-processing
# ============================================================
add_heading('3. Pre-processing Settings')

add_bullet('Frame size: ', '640 x 480 - all frames resized to same size for consistency')
add_bullet('Depth normalization: ', '(0.0, 1.0) - MiDaS raw output normalized. 0.0=close, 1.0=far')

add_heading('Median Filter (kernel size = 5)', level=2)
add_bullet('NOT a real kernel - just a 5x5 sliding window')
add_bullet('Collects all 25 values, sorts them, picks the MIDDLE value')
add_bullet('Removes noise (random wrong pixels) from depth map')
add_bullet('Better than average: Average of [50,50,50,250,50] = 90 (bad). Median = 50 (good)')
add_bullet('Preserves edges while removing outliers')

# ============================================================
# 4. Pose Estimation
# ============================================================
add_heading('4. Pose Estimation (MediaPipe)')

add_heading('Model Complexity', level=2)
add_table(
    ['Value', 'Name', 'Speed', 'Accuracy'],
    [
        ['0', 'Lite', 'Fastest', 'Basic'],
        ['1 (ours)', 'Full', 'Medium', 'Good'],
        ['2', 'Heavy', 'Slowest', 'Best'],
    ]
)

add_heading('Confidence Thresholds = 0.5', level=2)
add_bullet('Detection: "Only detect a person if >=50% sure"')
add_bullet('Tracking: "Keep tracking if >=50% sure it is the same person"')

add_heading('33 Landmarks', level=2)
add_para('Each body part has ONE point with a unique ID (0-32):')
add_bullet('0=nose, 7=left ear, 8=right ear')
add_bullet('11=left shoulder, 12=right shoulder')
add_bullet('23=left hip, 24=right hip')
add_bullet('25=left knee, 26=right knee')
add_bullet('27=left ankle, 28=right ankle')
add_bullet('Plus eyes, mouth, wrists, fingers, feet = 33 total')

add_heading('4 Values Per Landmark: (x, y, z, visibility)', level=2)
add_bullet('x', ' = horizontal position (0=left, 1=right)')
add_bullet('y', ' = vertical position (0=top, 1=bottom)')
add_bullet('z', ' = relative depth WITHIN body (arm in front vs behind, relative to hip center)')
add_bullet('visibility', ' = detection confidence (0=hidden, 1=clearly visible)')

add_heading('9 KEY_JOINTS (most important for fall detection)', level=2)
add_para('nose, left/right shoulder, left/right hip, left/right knee, left/right ankle')

# ============================================================
# 5. Pose z vs MiDaS Depth
# ============================================================
add_heading('5. Pose z vs MiDaS Depth - Why Both?')

add_table(
    ['', 'Pose z', 'MiDaS Depth'],
    [
        ['Reference', 'Body-relative (hip = 0)', 'Camera-relative'],
        ['Tells you', 'Which limb is in front/behind', 'How far person is from camera'],
        ['Example', '"Hand is in front of hip"', '"Person is 2.5 meters away"'],
        ['Falls toward camera', 'z does not change much', 'Depth decreases -> detected!'],
    ]
)
add_para('Both give DIFFERENT information. Combining them = better fall detection.')

# ============================================================
# 6. Feature Vector
# ============================================================
add_heading('6. Feature Vector (177 Dimensions)')

add_heading('Formula', level=2)
add_para('FEATURE_DIM = 33x4 + 33 + 3 + 9 = 177', bold=True)

add_table(
    ['Part', 'Count', 'What It Is'],
    [
        ['Pose landmarks (x,y,z,vis x 33)', '132', 'Body posture'],
        ['Depth at each joint', '33', 'Distance from camera per joint'],
        ['Bounding box aspect ratio', '1', 'Body shape (tall vs wide)'],
        ['Distance to floor', '1', 'How high hips are from ground'],
        ['Center of mass height', '1', 'Average height of ALL joints'],
        ['Vertical velocity (9 key joints)', '9', 'Speed of downward movement'],
        ['TOTAL', '177', ''],
    ]
)

add_heading('Bounding Box Aspect Ratio', level=2)
add_bullet('width / height of a box drawn around the person')
add_bullet('Standing: ratio ~ 0.4 (tall, narrow)')
add_bullet('Fallen: ratio ~ 3.0+ (wide, short - lying flat)')

add_heading('Distance to Floor', level=2)
add_bullet('1.0 - average_hip_y')
add_bullet('Standing: ~0.5 (hips at mid-frame)')
add_bullet('Fallen: ~0.15 (hips near ground)')

add_heading('Center of Mass Height', level=2)
add_bullet('Average y value of ALL 33 joints')
add_bullet('Different from floor distance: floor distance checks only hips, center of mass checks ENTIRE body')
add_bullet('Helps distinguish bending down (head still high) from falling (everything low)')

add_heading('Vertical Velocity', level=2)
add_bullet('= current_frame.y - previous_frame.y for each key joint')
add_bullet('MOST IMPORTANT feature for fall detection!', '')
add_bullet('Fall: large positive values (fast downward movement)')
add_bullet('Normal: near-zero values')
add_bullet('Key insight: falling happens FAST, sitting happens SLOWLY - velocity captures this')

# ============================================================
# 7. Temporal Sequence
# ============================================================
add_heading('7. Temporal Sequence')

add_table(
    ['Setting', 'Value', 'Meaning'],
    [
        ['FPS', '30', 'Camera captures 30 frames per second'],
        ['SEQUENCE_LENGTH', '30', 'LSTM looks at 30 frames (= 1 second) per decision'],
        ['SEQUENCE_STRIDE', '10', 'Sliding window moves by 10 frames (overlapping samples)'],
    ]
)
add_bullet('Stride = 10 creates overlapping samples, so each fall is seen from multiple perspectives')

# ============================================================
# 8. LSTM Model
# ============================================================
add_heading('8. LSTM Model Settings')

add_table(
    ['Setting', 'Value', 'Meaning'],
    [
        ['Hidden size', '128', 'Memory capacity - 128 numbers of understanding'],
        ['Layers', '2', 'Two stacked LSTMs: basic -> complex patterns'],
        ['Dropout', '0.3', 'Randomly turns off 30% neurons to prevent memorization'],
        ['Bidirectional', 'True', 'Reads sequence forward AND backward'],
    ]
)
add_heading('Why Bidirectional?', level=2)
add_bullet('Forward: "I know what happened before each frame"')
add_bullet('Backward: "I also know what happens after each frame"')
add_bullet('Combined: much better at understanding the full motion pattern')

# ============================================================
# 9. Training Settings
# ============================================================
add_heading('9. Training Settings')

add_table(
    ['Setting', 'Value', 'Meaning'],
    [
        ['Batch size', '16', 'Process 16 samples at once per step'],
        ['Learning rate', '0.001', 'Size of adjustment steps'],
        ['Weight decay', '0.0001', 'Prevents overcomplication, keeps model simple'],
        ['Max epochs', '50', 'Maximum training rounds (stopped at 25)'],
        ['Early stopping', '10', 'Stop if no improvement for 10 epochs'],
        ['LR scheduler patience', '5', 'Cut LR in half after 5 epochs no improvement'],
        ['LR scheduler factor', '0.5', 'Multiply learning rate by 0.5'],
        ['Train/Val/Test split', '70/15/15', '70% learning, 15% progress check, 15% final test'],
        ['Random seed', '42', 'Reproducibility (tradition from Hitchhiker\'s Guide)'],
    ]
)

add_heading('Classification Settings', level=2)
add_bullet('NUM_CLASSES = 1', ' - single output (0.0 to 1.0) for binary classification')
add_bullet('FALL_THRESHOLD = 0.5', ' - if output >= 0.5 then FALL, else Normal')
add_bullet('Can adjust: 0.3 for safety-critical, 0.7 for low false alarms')

# ============================================================
# 10. Inference
# ============================================================
add_heading('10. Inference Settings')

add_heading('Fall Confirmation Window', level=2)
add_para('Person goes down -> "FALL SUSPECTED" (orange, wait 3s) -> Still down? -> "FALL CONFIRMED!" (red)')
add_para('Person gets back up within 3s -> "Normal" (cancelled, no alert)')

add_table(
    ['Setting', 'Value', 'Meaning'],
    [
        ['FALL_CONFIRMATION_SECONDS', '3', 'Wait 3 seconds before confirming fall'],
        ['RECOVERY_THRESHOLD', '0.3', 'If confidence drops below 30%, person recovered'],
        ['ALERT_COOLDOWN_SECONDS', '5', 'After confirmed fall, wait 5s before next alert'],
    ]
)
add_bullet('Gap between thresholds (0.3 to 0.5) prevents flickering/false triggers')

# ============================================================
# 11. Results
# ============================================================
add_heading('11. Results Explained')

add_para('Test set: 21 samples (7 falls, 14 normal)', bold=True)

add_table(
    ['Metric', 'Value', 'Meaning'],
    [
        ['Accuracy', '100%', 'All 21 predictions correct'],
        ['Precision', '100%', 'When it says FALL, always right'],
        ['Sensitivity/Recall', '100%', 'Catches ALL falls'],
        ['Specificity', '100%', 'Correctly identifies ALL normal activities'],
        ['F1 Score', '100%', 'Perfect balance of precision and recall'],
        ['False Alarm Rate', '0%', 'Never cries wolf'],
        ['AUC-ROC', '1.0', 'Overall quality = perfect'],
    ]
)

add_heading('Confusion Matrix', level=2)
add_bullet('True Positives (TP) = 7', ' - was fall, said fall')
add_bullet('True Negatives (TN) = 14', ' - was normal, said normal')
add_bullet('False Positives (FP) = 0', ' - no false alarms')
add_bullet('False Negatives (FN) = 0', ' - no missed falls')

add_heading('Why 100% is Misleading', level=2)
add_bullet('Small test set (only 21 samples) - model may have memorized')
add_bullet('Expected realistic accuracy on full dataset: 90-95%')
add_bullet('Need more data and cross-validation for trustworthy results')

# ============================================================
# 12. Limitations
# ============================================================
add_heading('12. Known Limitations & Professor Answers')

add_heading('Exercise Detection', level=2)
add_bullet('Current model may confuse push-ups/sit-ups with falls (never trained on exercise videos)')
add_bullet('Solution: add NTU RGB+D exercise data in weeks 4-5')

add_heading('Live Camera Performance', level=2)
add_bullet('Runs at ~2-5 FPS on CPU (MiDaS + pose per frame)')
add_bullet('GPU would give ~15-30 FPS')

add_heading('Why MiDaS instead of real depth camera?', level=2)
add_bullet('Eliminates need for expensive hardware (Kinect, RealSense)')
add_bullet('Works with ANY regular webcam - key innovation of this project')

add_heading('Is it a replica?', level=2)
add_bullet('NO - unique combination of MiDaS + 33 pose landmarks + 177-dim features + Bidirectional LSTM with attention')
add_bullet('Uses standard tools combined in an original way')
add_bullet('This is how real research works - build on existing tools, combine them innovatively')

# ============================================================
# 13. Files
# ============================================================
add_heading('13. Project Files Summary')

add_table(
    ['File', 'Purpose'],
    [
        ['config.py', 'All settings in one place'],
        ['src/models/depth_estimator.py', 'MiDaS depth map generation'],
        ['src/models/pose_estimator.py', '33-joint skeleton detection'],
        ['src/models/feature_extractor.py', 'Creates 177-dim feature vector'],
        ['src/models/fall_detector.py', 'LSTM classifier with attention'],
        ['src/data/download_dataset.py', 'Downloads URFD dataset'],
        ['src/data/preprocess.py', 'Processes videos into features'],
        ['src/data/dataset.py', 'PyTorch data loader with sliding window'],
        ['src/train.py', 'Training loop with early stopping'],
        ['src/evaluate.py', 'Computes metrics, generates plots'],
        ['src/inference.py', 'Live camera + video file detection'],
        ['src/utils/visualization.py', 'Plotting helpers'],
        ['src/utils/helpers.py', 'Common utilities'],
    ]
)

add_heading('Pipeline Flow', level=2)
add_para('Frame -> MiDaS (depth) -> MediaPipe (pose) -> Feature Extractor (177 numbers) -> LSTM (30 frames) -> Confirmation Window (3s) -> FALL or NORMAL', bold=True)


# ---- Save ----
output_path = Path(__file__).parent.parent / "docs" / "study_notes.docx"
doc.save(str(output_path))
print(f"Saved: {output_path}")
